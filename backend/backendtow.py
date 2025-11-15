import os
import logging
import time
import tempfile
import hashlib
import json
import mimetypes

import google.generativeai as genai
import whisper
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from sentence_transformers import CrossEncoder
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from PIL import Image
from transformers import BlipProcessor, BlipForConditionalGeneration
import tiktoken
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
# from nltk.tokenize import sent_tokenize  # Lazy load
from .config import AUTO_PERSIST_STRUCTURED
from .data_extractor import DataExtractor, DataProcessor
from .models import db as sqldb

# === CONFIGURATION ===
INDEX_PATH = os.path.join(os.getcwd(), "index/arx_faiss")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "")
CACHE_DIR = os.path.join(os.getcwd(), "cache")
os.makedirs(CACHE_DIR, exist_ok=True)
os.makedirs(os.path.dirname(INDEX_PATH), exist_ok=True)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Initialisation lazy des modèles (chargement à la demande)
genai.configure(api_key=GOOGLE_API_KEY)
model = None
whisper_model = None

# Embeddings et modèles
embeddings = None
cross_encoder = None
db = None

# Tokenizer pour compter précisément les tokens
tokenizer = tiktoken.get_encoding("gpt2")

# === BLIP IMAGE CAPTIONING ===
blip_processor = None
blip_model = None

# Fonctions lazy loading
def get_model():
    global model
    if model is None:
        logging.info("Chargement du modèle Gemini...")
        model = genai.GenerativeModel("models/gemini-2.5-pro")
    return model

def get_whisper_model():
    global whisper_model
    if whisper_model is None:
        logging.info("Chargement du modèle Whisper...")
        whisper_model = whisper.load_model("base")
    return whisper_model

def get_embeddings():
    global embeddings
    if embeddings is None:
        logging.info("Chargement des embeddings...")
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    return embeddings

def get_cross_encoder():
    global cross_encoder
    if cross_encoder is None:
        logging.info("Chargement du cross-encoder...")
        cross_encoder = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")
    return cross_encoder

def get_blip_models():
    global blip_processor, blip_model
    if blip_processor is None:
        logging.info("Chargement des modèles BLIP...")
        blip_processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
        blip_model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
    return blip_processor, blip_model

def describe_image_with_blip(image_path):
    try:
        blip_processor, blip_model = get_blip_models()
        image = Image.open(image_path).convert("RGB")
        inputs = blip_processor(images=image, return_tensors="pt")
        out = blip_model.generate(**inputs)
        caption = blip_processor.decode(out[0], skip_special_tokens=True)
        return caption
    except Exception as e:
        return f"[Erreur lors de la description de l'image : {str(e)}]"

# === UTILS ===

def chunk_text_semantically(text, max_tokens=500, overlap_tokens=100):
    from nltk.tokenize import sent_tokenize
    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = []
    current_len = 0

    for sentence in sentences:
        sent_tokens = count_tokens(sentence)

        if current_len + sent_tokens > max_tokens:
            chunks.append(" ".join(current_chunk))
            # Overlap sémantique
            overlap = []
            token_sum = 0
            for sent in reversed(current_chunk):
                token_sum += count_tokens(sent)
                overlap.insert(0, sent)
                if token_sum >= overlap_tokens:
                    break
            current_chunk = overlap
            current_len = sum(count_tokens(s) for s in current_chunk)

        current_chunk.append(sentence)
        current_len += sent_tokens

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks

def generate_export_file(data, format="txt"):
    answer = data.get("answer", "Aucune réponse")
    context = data.get("context", [])

    if format == "txt":
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8")
        temp.write("Réponse générée :\n")
        temp.write(answer + "\n\n")
        if context:
            temp.write("Contexte utilisé :\n")
            for i, doc in enumerate(context):
                temp.write(f"Chunk {i+1} : {doc.page_content}\n---\n")
        temp.close()
        return temp.name

    elif format == "docx":
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc = DocxDocument()
        doc.add_heading("Réponse générée", 0)
        doc.add_paragraph(answer)
        if context:
            doc.add_heading("Contexte utilisé", level=1)
            for i, doc_chunk in enumerate(context):
                doc.add_paragraph(f"Chunk {i+1} :", style='Heading2')
                doc.add_paragraph(doc_chunk.page_content)
        doc.save(temp.name)
        return temp.name

    elif format == "pdf":
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        doc = SimpleDocTemplate(temp.name)
        styles = getSampleStyleSheet()
        flowables = [
            Paragraph("Réponse générée", styles["Heading1"]),
            Paragraph(answer, styles["Normal"]),
        ]
        if context:
            flowables.append(Paragraph("Contexte utilisé", styles["Heading2"]))
            for i, doc_chunk in enumerate(context):
                flowables.append(Paragraph(f"Chunk {i+1} :", styles["Heading3"]))
                flowables.append(Paragraph(doc_chunk.page_content, styles["Normal"]))
        doc.build(flowables)
        return temp.name

    else:
        raise ValueError("Format non pris en charge")

def get_file_hash(file_bytes):
    return hashlib.md5(file_bytes).hexdigest()

def get_title_from_filename(filename):
    return os.path.splitext(os.path.basename(filename))[0]

def count_tokens(text):
    return len(tokenizer.encode(text))

def truncate_text_by_tokens(text, max_tokens):
    tokens = tokenizer.encode(text)
    if len(tokens) <= max_tokens:
        return text
    return tokenizer.decode(tokens[:max_tokens])

def summarize_text(text, max_chars=500):
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "..."

def summarize_chat_history(chat_history, max_chars=1000):
    full_text = ""
    for turn in chat_history:
        full_text += f"Utilisateur : {turn['user']}\nAssistant : {turn['assistant']}\n"
    return summarize_text(full_text, max_chars=max_chars)

from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

def evaluate_answer_quality(answer: str, context_docs: list, model_name="sentence-transformers/all-MiniLM-L6-v2"):
    try:
        from sentence_transformers import SentenceTransformer
        sbert_model = SentenceTransformer(model_name)

        # Embedding de la réponse
        answer_embedding = sbert_model.encode([answer], convert_to_tensor=True)

        # Embedding des chunks contextuels
        doc_texts = [doc.page_content for doc in context_docs]
        doc_embeddings = sbert_model.encode(doc_texts, convert_to_tensor=True)

        # Similarité cosine entre la réponse et chaque chunk
        similarities = cosine_similarity(answer_embedding, doc_embeddings)[0]

        # Score final : moyenne des similarités
        mean_score = np.mean(similarities)
        return round(mean_score, 4)  

    except Exception as e:
        logging.warning(f"Erreur évaluation qualité : {e}")
        return None

# === FAISS INDEX ===

def load_faiss_index():
    global db
    try:
        emb = get_embeddings()
        db = FAISS.load_local(INDEX_PATH, emb, allow_dangerous_deserialization=True)
        logging.info(" Index FAISS chargé.")
    except Exception as e:
        logging.warning(f" Index non trouvé ou invalide, création d'un index vide : {e}")
        # Créer un document placeholder pour initialiser l'index
        emb = get_embeddings()
        placeholder = Document(page_content="placeholder", metadata={"source": "init"})
        db = FAISS.from_documents([placeholder], emb)
        # Supprimer le document placeholder
        db.delete([db.index_to_docstore_id[0]])
        db.save_local(INDEX_PATH)
        logging.info(" Index FAISS vide créé.")

def reload_faiss_index():
    logging.info(" Rechargement de l’index FAISS...")
    load_faiss_index()

def reset_faiss_index():
    global db
    logging.info(" Réinitialisation de l'index FAISS...")
    emb = get_embeddings()
    placeholder = Document(page_content="placeholder", metadata={"source": "init"})
    db = FAISS.from_documents([placeholder], emb)
    db.delete([db.index_to_docstore_id[0]])
    db.save_local(INDEX_PATH)
    logging.info(" Index FAISS réinitialisé.")

def get_existing_document_ids():
    try:
        return set(doc.metadata.get("document_id") for doc in db.similarity_search("", k=1000) if doc.metadata.get("document_id"))
    except Exception as e:
        logging.warning(f"Erreur récupération des document_id : {e}")
        return set()

def add_document_to_index(text, metadata=None):
    global db
    try:
        if not text.strip():
            logging.warning("Texte vide, rien à indexer.")
            return False

        document_id = metadata.get("document_id") if metadata else None
        existing_ids = get_existing_document_ids()

        if document_id and document_id in existing_ids:
            logging.info(f" Document déjà indexé : {document_id}")
            return False  

        chunks = chunk_text_semantically(text, max_tokens=500, overlap_tokens=100)

        docs = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy() if metadata else {}
            chunk_metadata.update({
                "chunk_index": i,
                "chunk_length": len(chunk),
                "title": metadata.get("title", "Sans titre") if metadata else "Sans titre"
            })

            docs.append(Document(page_content=chunk, metadata=chunk_metadata))

        db.add_documents(docs)
        db.save_local(INDEX_PATH)
        logging.info(f" {len(docs)} chunks ajoutés avec métadonnées enrichies.")
        return True
    except Exception as e:
        logging.error(f"Erreur ajout document à l'index : {e}")
        return False

# === STRUCTURED DATA AUTO-PERSISTENCE ===

def _auto_persist_structured(text: str, source_name: str, source_type: str = "FILE"):
    """Analyse le texte avec Gemini et persiste en BD si activé.

    Cette opération est best-effort: en cas d'erreur (quota, réseau), on log et on continue.
    """
    try:
        if not AUTO_PERSIST_STRUCTURED:
            return
        if not text or len(text.strip()) < 50:
            # Evite les très petits contenus
            return
        if not GOOGLE_API_KEY:
            logging.info("AUTO_PERSIST_STRUCTURED actif mais GOOGLE_API_KEY manquant; skip")
            return

        extractor = DataExtractor(GOOGLE_API_KEY)
        structured = extractor.parse_ingredients_and_products(text)
        has_payload = any([
            structured.get("products"),
            structured.get("ingredients_info"),
            structured.get("incompatibilities"),
        ])
        if has_payload:
            # On utilise le nom du fichier comme 'source_file'
            DataProcessor.process_extraction(structured, sqldb, source_name)
            logging.info("Données structurées persistées automatiquement")
    except Exception as e:
        logging.warning(f"Auto-persist structuré ignoré: {e}")

# === EXTRACTION ===

def extract_text(file):
    try:
        file_bytes = file.read()
        file.seek(0)
        file_hash = get_file_hash(file_bytes)
        cache_path = os.path.join(CACHE_DIR, f"{file_hash}_text.json")
        if os.path.exists(cache_path):
            logging.info("Chargement texte extrait en cache")
            return json.load(open(cache_path, "r", encoding="utf-8"))["text"]

        ext = os.path.splitext(file.filename)[1].lower()
        mime_type, _ = mimetypes.guess_type(file.filename)
        text = ""

        if ext == ".pdf":
            text = "\n".join([p.extract_text() or "" for p in PdfReader(file).pages])
        elif ext == ".docx":
            text = "\n".join([p.text for p in DocxDocument(file).paragraphs])
        elif ext in [".png", ".jpg", ".jpeg"]:
            file.seek(0)
            text = describe_image_with_blip(file)
        elif "audio" in (mime_type or "") or ext in [".mp3", ".wav", ".m4a"]:
            text = transcribe_audio(file)
        else:
            file.seek(0)
            text = file.read().decode("utf-8", errors="ignore")

        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump({"text": text}, f, ensure_ascii=False, indent=2)

        file.seek(0)
        return text
    except Exception as e:
        logging.error(f"Erreur extraction : {e}")
        return f"Erreur extraction : {e}"

def transcribe_audio(audio_input):
    try:
        file_bytes = audio_input.read()
        audio_input.seek(0)
        file_hash = get_file_hash(file_bytes)
        cache_path = os.path.join(CACHE_DIR, f"{file_hash}_asr.json")
        if os.path.exists(cache_path):
            logging.info("Chargement transcription en cache")
            return json.load(open(cache_path, "r", encoding="utf-8"))["text"]
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        asr_model = get_whisper_model()
        result = asr_model.transcribe(tmp_path)
        os.remove(tmp_path)
        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        return result["text"]
    except Exception as e:
        logging.error(f"Erreur transcription : {e}")
        return f"Erreur transcription : {e}"

# === RERANKING ===

def rerank_documents(query, docs, top_k=4):
    if not docs:
        return []
    pairs = [(query, doc.page_content) for doc in docs]
    encoder = get_cross_encoder()
    scores = encoder.predict(pairs)
    scored_docs = sorted(zip(scores, docs), key=lambda x: x[0], reverse=True)
    return [doc for _, doc in scored_docs[:top_k]]

# === RAG FUSION MULTI-DOCS ===

def rag_fusion_multi_docs(
    query, 
    chat_history=None, 
    k=6, 
    max_context_tokens=1500, 
    max_history_tokens=1000, 
    nb_messages=5, 
    retries=3
):
    if db is None:
        logging.error("Index FAISS non chargé")
        return " Index non chargé.", []

    if chat_history is None:
        chat_history = []

    try:
        # Recherche dans FAISS
        retrieved_docs = db.similarity_search(query, k=k)
        # Reranking des documents les plus pertinents
        docs = rerank_documents(query, retrieved_docs, top_k=k)
    except Exception as e:
        logging.error(f"Erreur recherche documentaire : {e}")
        return f"Erreur recherche documentaire : {e}", []

    context_text = ""
    context_token_count = 0
    context_docs = []

    # Construction du contexte fusionné en respectant max tokens
    for doc in docs:
        doc_tokens = count_tokens(doc.page_content)
        if context_token_count + doc_tokens > max_context_tokens:
            break
        source_info = doc.metadata.get("title", "Document inconnu")
        context_text += f"[Source: {source_info}]\n{doc.page_content}\n\n"
        context_token_count += doc_tokens
        context_docs.append(doc)

    # Préparation historique résumé en respectant max tokens
    summarized_history = ""
    history_token_count = 0
    for turn in chat_history[-nb_messages:]:
        turn_text = f"Utilisateur : {turn['user']}\nAssistant : {turn['assistant']}\n"
        turn_tokens = count_tokens(turn_text)
        if history_token_count + turn_tokens > max_history_tokens:
            break
        summarized_history += turn_text
        history_token_count += turn_tokens

    summarized_history = truncate_text_by_tokens(summarized_history, max_history_tokens)

    prompt = f"""
Tu es un assistant IA expert. Voici une question d'utilisateur, des extraits documentaires provenant de plusieurs documents/fichiers, ainsi qu'un historique résumé du dialogue.

### ❓ Question :
{query}

### 📚 Contexte documentaire multi-docs fusionné :
{context_text}

### 💬 Historique résumé :
{summarized_history}

### ✍️ Réponse :
""".strip()

    for attempt in range(retries):
        try:
            llm_model = get_model()
            response = llm_model.generate_content(prompt)
            #full_answer = response.text.strip()
            if response.candidates and response.candidates[0].content.parts:
                full_answer = "".join(
                    [part.text for part in response.candidates[0].content.parts if hasattr(part, "text")]
                ).strip()
            else:
                full_answer = " Désolé, je n'ai pas pu générer de réponse."


            # Résumé automatique de la réponse pour simplicité
            summary_prompt = f"""
Résume ce texte en 1 à 2 phrases simples et claires, comme une synthèse pour un utilisateur non-expert.

Texte :
{full_answer}

Résumé :
""".strip()
            summary_response = llm_model.generate_content(summary_prompt)
            answer_summary = summary_response.text.strip()

            # Ne pas afficher les chunks utilisés dans la réponse
            final_response = f"{full_answer}\n\n📄 Résumé : {answer_summary}"

            return final_response, context_docs

        except Exception as e:
            logging.warning(f"Tentative {attempt+1} échouée : {e}")
            time.sleep(2)

    return "❌ Réponse impossible.", context_docs



# === RAG DIRECT PROMPT ===

def rag_direct_prompt(query, chat_history=None, nb_messages=5, retries=3):
    if chat_history is None:
        chat_history = []
    prompt = build_prompt_with_context(chat_history[-nb_messages:], query)
    for attempt in range(retries):
        try:
            llm_model = get_model()
            response = llm_model.generate_content(prompt)
            return response.text.strip()
        except Exception as e:
            logging.warning(f"Tentative {attempt+1} échouée : {e}")
            time.sleep(2)
    return "❌ Réponse impossible."

# === CONSTRUCTION PROMPT ===

def build_prompt_with_context(chat_history, query):
    summarized_history = ""
    for turn in chat_history:
        summarized_history += f"Utilisateur : {turn['user']}\nAssistant : {turn['assistant']}\n"
    prompt = f"""
Tu es un assistant IA expert. Voici une question d'utilisateur et un historique résumé du dialogue.

### ❓ Question :
{query}

### 💬 Historique résumé :
{summarized_history}

### ✍️ Réponse :
""".strip()
    return prompt

# === PROCESS QUESTION ===

def process_question(query, use_rag=True, chat_history=None, nb_messages=5):
    if use_rag:
        # Utiliser la fusion multi-docs ici
        answer, _ = rag_fusion_multi_docs(query, chat_history, nb_messages=nb_messages)
    else:
        answer = rag_direct_prompt(query, chat_history, nb_messages=nb_messages)
    return answer

# === HANDLE UPLOADED FILE ===

def handle_uploaded_file(file, question=None, chat_history=None, use_rag=True, nb_messages=5):
    text = extract_text(file)
    if not text or "Erreur" in text:
        return text

    file_bytes = file.read()
    file.seek(0)
    doc_id = get_file_hash(file_bytes)
    title = get_title_from_filename(file.filename)

    metadata = {
        "document_id": doc_id,
        "source": file.filename,
        "title": title
    }

    add_document_to_index(text, metadata=metadata)
    # Persistance automatique des données structurées (best-effort)
    try:
        _auto_persist_structured(text, source_name=file.filename, source_type=os.path.splitext(file.filename)[1].lower())
    except Exception:
        pass

    if question:
        prompt = f"{question.strip()}\n\nContenu du fichier :\n{text.strip()}"
    else:
        prompt = text.strip()

    return process_question(prompt, use_rag=use_rag, chat_history=chat_history, nb_messages=nb_messages)

def handle_multiple_uploaded_files(files, question=None, chat_history=None, use_rag=True, nb_messages=5):
    """
    Traite plusieurs fichiers uploadés et génère une réponse RAG combinée.
    """
    all_text = ""
    for file in files:
        text = extract_text(file)
        if not text or "Erreur" in text:
            continue  # Ignore les fichiers avec erreur
        file_bytes = file.read()
        file.seek(0)
        doc_id = get_file_hash(file_bytes)
        title = get_title_from_filename(file.filename)
        ext = os.path.splitext(file.filename)[1].lower()
        metadata = {
            "document_id": doc_id,
            "source": file.filename,
            "title": title
        }
        add_document_to_index(text, metadata=metadata)
        # Persistance automatique par fichier (best-effort)
        try:
            _auto_persist_structured(text, source_name=file.filename, source_type=ext)
        except Exception:
            pass
        all_text += f"\n\n### Fichier : {file.filename} ###\n{text.strip()}"

    if not all_text.strip():
        return "Aucun contenu exploitable trouvé dans les fichiers."

    if question:
        prompt = f"{question.strip()}\n\nContenu combiné des fichiers :\n{all_text.strip()}"
    else:
        prompt = all_text.strip()

    return process_question(prompt, use_rag=use_rag, chat_history=chat_history, nb_messages=nb_messages)

# === INIT ===
load_faiss_index()
